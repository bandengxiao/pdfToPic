from docx import Document
import os

# document = Document()
# document.add_heading('工具使用方法', 0)
#
# p = document.add_paragraph()
# p.add_run('1、pdf地址处填写生成pdf文件夹路径\n2、图片文件地址处填写图片文件夹路径\n'
#           '3、pdf生成图片时在图片地址内按pdf文件生成目标文件夹，如A.pdf会生成A文件夹，A文件夹内存放A.pdf生成的所有图片文件\n'
#           '4、图片生成pdf时会遍历路径下的所有二级文件夹，每个文件夹里的所有图片会生成一个pdf文件'
#           ).bold = True
#
#
# document.save('demo.docx')

test=os.path.dirname(r'C:\Users\92410\Desktop\pdf\新建文件夹\test.py')
print(test)
print(os.path.exists(test))