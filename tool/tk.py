import tkinter as tk
from tkinter import messagebox
from tkinter import ttk,Menu,filedialog
import os
from tool.pdf_toPic import pdf_toPic
from tool.pic_toPDF import picToPDF
from docx import Document


class tk_main():

    def __init__(self):
        self.root = tk.Tk()
        self.root.title('pdf、图片转换工具')
        # self.root.geometry("500x200+800+300")
        # self.root.resizable(False,False)
        self.root.geometry("500x200+800+300")
        self.menubar = Menu(self.root)  # 主菜单实例
        self.root.config(menu=self.menubar)  # 显示菜单,将root根窗口的主菜单设置为menu

        # 创建主菜单实例
        self.menubar = Menu(self.root)
        # 显示菜单,将root根窗口的主菜单设置为menu
        self.root.config(menu=self.menubar)

        self.interface()

    def interface(self):
        self.Label0 = tk.Label(self.root, text="PDF文件地址 :")
        self.Label0.place(x=1, y=50)

        self.Label0 = tk.Label(self.root, text="图片文件地址 :")
        self.Label0.place(x=1, y=110)

        self.Label0 = tk.Label(self.root, text="请选择操作项 :")
        self.Label0.place(x=1, y=10)

        self.pdf_W=tk.Text(self.root, width=50, height=2,undo = True)
        self.pdf_W.insert("insert","输入PDF地址...")
        self.pdf_W.config(state='disabled')
        self.pdf_W.place(x=100, y=50)

        self.pic_W = tk.Text(self.root, width=50, height=2,undo = True)
        self.pic_W.insert("insert", "输入图片地址...")
        self.pic_W.config(state='disabled')
        self.pic_W.place(x=100, y=110)

        self.value = tk.StringVar()
        self.value.set('2')  # 默认值
        values = ['请选择','pdf转图片', '图片转pdf']
        self.combobox = ttk.Combobox(
            master=self.root,  # 父容器
            height=5,  # 高度,下拉显示的条目数量
            width=20,  # 宽度
            state='',  # 设置状态 normal(可选可输入)、readonly(只可选)、 disabled(禁止输入选择)
            cursor='arrow',  # 鼠标移动时样式 arrow, circle, cross, plus...
            font=('', 10),  # 字体、字号
            textvariable=self.value,  # 通过StringVar设置可改变的值
            values=values,  # 设置下拉框的选项
        )
        self.combobox.current(0)
        self.combobox.place(x=100, y=12)
        # 绑定事件,下拉列表框被选中时，绑定pick()函数
        self.combobox.bind("<<ComboboxSelected>>", self.pick)
        self.Button0=tk.Button(self.root,)

        self.Button0=tk.Button(self.root,text='开始转换', command=self.event)
        self.Button0.place(x=210, y=180)

        self.Button1=tk.Button(self.root, bitmap='question',activebackground='#4A766E',command=self.help)
        self.Button1.place(x=460, y=6)


    def help(self):

        try:
            # 去掉title='保存为'，标题默认名称“另存为”
            file_name = filedialog.asksaveasfilename(title='导出帮助文档到...', filetypes=[("docx", ".docx")],
                                                     defaultextension=".*")
            # 文件添加数据并保存
            document = Document()
            document.add_heading('工具使用方法', 0)

            p = document.add_paragraph()
            p.add_run('1、pdf地址处填写生成pdf文件夹路径\n\n'
                      '2、图片文件地址处填写图片文件夹路径\n\n'
                      '3、pdf生成图片时在图片地址内按pdf文件生成目标文件夹，如A.pdf会生成A文件夹，A文件夹内存放A.pdf生成的所有图片文件\n\n'
                      '4、图片生成pdf时会遍历路径下的所有二级文件夹，每个文件夹里的所有图片会生成一个pdf文件'
                      ).bold = True
            if os.path.exists(os.path.dirname(file_name)):
                document.save(file_name)
                tk.messagebox.showinfo("运行结果", "帮助文档导出成功！")
            else:
                pass
        except Exception as e:
            tk.messagebox.showinfo("error",e)

    def pick(self,*args):
        select=self.combobox.get()
        if select != '请选择':
            self.pdf_W.config(state='normal')
            self.pdf_W.delete(1.0, "end")
            self.pic_W.config(state='normal')
            self.pic_W.delete(1.0, "end")
        else:

            self.pdf_W.insert("insert", "输入PDF地址...")
            self.pdf_W.config(state='disabled')

            self.pic_W.insert("insert", "输入图片地址...")
            self.pic_W.config(state='disabled')

    def event(self):
        self.pdf_address=self.pdf_W.get('0.0', 'end')
        self.pic_address=self.pic_W.get('0.0', 'end')
        func=self.combobox.get()
        if '请选择' == func:
            messagebox.showinfo('error','未选择操作项！')
        else:
            if 'pdf转图片' == func:
                try:
                    pdfToPic = pdf_toPic()
                    pdfToPic.tmp = str(self.pdf_address).strip()
                    pdfToPic.export_file = str(self.pic_address).strip()
                    self.pdf_dir = [i for i in os.listdir(str(self.pdf_address).strip()) if
                                    os.path.splitext(i)[-1] == ".pdf"]
                    pdfToPic.pdf_dir = self.pdf_dir
                    result = pdfToPic.main()
                    messagebox.showinfo('运行结果', result)
                except Exception as e:
                    messagebox.showinfo('error', e)
            elif '图片转pdf' == func:
                try:
                    ToPDF = picToPDF()
                    MessageResult=ToPDF.pic_to_Pdf(pathDir=str(self.pic_address).strip(),saveDir=str(self.pdf_address).strip())
                    messagebox.showinfo("运行结果",MessageResult)
                except Exception as e:
                    messagebox.showinfo("error",e)



if __name__ == '__main__':
    tk_window = tk_main()
    tk_window.root.mainloop()